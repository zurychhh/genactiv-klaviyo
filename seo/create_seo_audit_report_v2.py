#!/usr/bin/env python3
"""
Generate SEO Audit Report for GenActiv.pl - VERIFIED VERSION
Diplomatic wording, agency collaboration friendly
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, data, header_color="0066CC", alt_row_color="F5F5F5"):
    """Create a styled table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

            if i == 0:  # Header row
                set_cell_shading(cell, header_color)
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif i % 2 == 0:  # Alternating rows
                set_cell_shading(cell, alt_row_color)

    return table

def create_report():
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== TITLE PAGE =====
    title = doc.add_heading('AUDYT SEO GENACTIV.PL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 102, 204)

    subtitle = doc.add_paragraph('Raport techniczny z weryfikacją stanu faktycznego')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Data audytu: {datetime.now().strftime("%d.%m.%Y")}\n').bold = True
    info.add_run('Metodologia: Analiza kodu źródłowego (curl), MCP Shopify, sitemap XML\n')
    info.add_run('Zakres: Strona główna, 10 produktów, 10 kolekcji, 4 strony statyczne')

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading('Spis treści', level=1)
    toc_items = [
        '1. Podsumowanie wykonawcze',
        '2. Elementy SEO - stan zweryfikowany',
        '3. Obszary do optymalizacji',
        '4. Analiza źródeł ruchu',
        '5. Rekomendacje i priorytety',
        '6. Podział zadań operacyjnych'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # ===== SECTION 1: EXECUTIVE SUMMARY =====
    doc.add_heading('1. Podsumowanie wykonawcze', level=1)

    summary_intro = doc.add_paragraph()
    summary_intro.add_run('Ocena ogólna: ').bold = True
    summary_intro.add_run('Strona genactiv.pl posiada solidne fundamenty techniczne SEO. ')
    summary_intro.add_run('Większość kluczowych elementów jest poprawnie wdrożona, ')
    summary_intro.add_run('co stanowi dobrą bazę do dalszej optymalizacji.')

    doc.add_paragraph()

    # Score card
    score_data = [
        ['Obszar', 'Ocena', 'Status'],
        ['Technical SEO (meta, canonical, sitemap)', '8/10', 'Bardzo dobry'],
        ['Structured Data (Schema.org)', '8/10', 'Bardzo dobry'],
        ['Open Graph / Social', '9/10', 'Doskonały'],
        ['Content Coverage', '7/10', 'Dobry'],
        ['Źródła ruchu', '7/10', 'Dobry'],
    ]
    add_styled_table(doc, score_data, header_color="27AE60")

    doc.add_paragraph()

    key_findings = doc.add_paragraph()
    key_findings.add_run('Kluczowe ustalenia:\n').bold = True

    findings = [
        'Produkty mają kompletne meta descriptions, Open Graph i Schema.org Product',
        'System recenzji prawidłowo generuje AggregateRating w structured data',
        'Sitemap jest kompletny i zawiera daty modyfikacji (lastmod)',
        'Zidentyfikowano obszary do uzupełnienia: część kolekcji i strona FAQ'
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # ===== SECTION 2: VERIFIED SEO ELEMENTS =====
    doc.add_heading('2. Elementy SEO - stan zweryfikowany', level=1)

    doc.add_heading('2.1 Meta descriptions', level=2)

    meta_data = [
        ['Typ strony', 'Sprawdzono', 'Wynik', 'Uwagi'],
        ['Strona główna', '1', '1/1 OK', 'Poprawny, zawiera słowa kluczowe'],
        ['Produkty', '10 losowych', '10/10 OK', 'Wszystkie sprawdzone mają unikalne opisy'],
        ['Główne kolekcje', '10', '9/10 OK', '1 brak: produkty-dla-dzieci'],
        ['Strony statyczne', '4', '3/4 OK', '1 brak: kontakt'],
    ]
    add_styled_table(doc, meta_data, header_color="0066CC")

    doc.add_paragraph()
    meta_note = doc.add_paragraph()
    meta_note.add_run('Uwaga: ').bold = True
    meta_note.add_run('Zidentyfikowano 16 kolekcji bez meta descriptions (głównie kolekcje pomocnicze i sezonowe). ')
    meta_note.add_run('Lista w sekcji 3.')

    doc.add_heading('2.2 Structured Data (Schema.org)', level=2)

    schema_data = [
        ['Typ Schema', 'Obecność', 'Szczegóły'],
        ['Product', 'TAK', 'name, sku, price, priceCurrency, availability'],
        ['Offer', 'TAK', 'Zintegrowany z Product'],
        ['AggregateRating', 'TAK', 'ratingValue, ratingCount (dla produktów z recenzjami)'],
        ['Review', 'TAK', 'Indywidualne recenzje z ratingValue'],
        ['Organization', 'TAK', 'Na wszystkich stronach'],
        ['WebSite + SearchAction', 'TAK', 'Umożliwia sitelinks searchbox w Google'],
        ['BreadcrumbList', 'NIE', 'Breadcrumbs w HTML są, brak schema'],
        ['FAQPage', 'NIE', 'Strona FAQ istnieje, brak schema'],
    ]
    add_styled_table(doc, schema_data, header_color="0066CC")

    doc.add_heading('2.3 Open Graph', level=2)

    og_data = [
        ['Tag', 'Obecność', 'Przykład'],
        ['og:title', 'TAK', 'Colostrum Genactiv Kapsułki 120 kapsułek – Genactiv'],
        ['og:description', 'TAK', 'Kup Colostrum Genactiv Kapsułki...'],
        ['og:image', 'TAK', 'Zdjęcie produktu w wysokiej rozdzielczości'],
        ['og:price:amount', 'TAK', '189,00'],
        ['og:price:currency', 'TAK', 'PLN'],
        ['og:type', 'TAK', 'product / website'],
        ['og:url', 'TAK', 'Canonical URL'],
    ]
    add_styled_table(doc, og_data, header_color="0066CC")

    og_note = doc.add_paragraph()
    og_note.add_run('Wynik: ').bold = True
    og_note.add_run('10 tagów Open Graph na stronach produktowych - kompletna implementacja.')

    doc.add_heading('2.4 Elementy techniczne', level=2)

    tech_data = [
        ['Element', 'Status', 'Wartość'],
        ['Canonical URLs', 'OK', 'Poprawne na wszystkich sprawdzonych stronach'],
        ['Sitemap.xml', 'OK', 'Index + 5 child sitemaps (products, pages, collections, blogs, metaobject)'],
        ['Sitemap lastmod', 'OK', 'Daty aktualizacji obecne we wszystkich sitemapach'],
        ['Robots.txt', 'OK', 'Prawidłowa konfiguracja Shopify'],
        ['Viewport', 'OK', 'width=device-width,initial-scale=1'],
        ['Charset', 'OK', 'UTF-8'],
        ['Lang', 'OK', 'pl'],
        ['HTTPS', 'OK', 'Wymuszony'],
    ]
    add_styled_table(doc, tech_data, header_color="27AE60")

    doc.add_page_break()

    # ===== SECTION 3: AREAS FOR OPTIMIZATION =====
    doc.add_heading('3. Obszary do optymalizacji', level=1)

    doc.add_paragraph('Poniżej szczegółowa lista elementów wymagających uzupełnienia:')

    doc.add_heading('3.1 Kolekcje bez meta descriptions (16)', level=2)

    coll_list = [
        'produkty-dla-dzieci', 'akcesoria', 'maseczki-z-colostrum', 'kremy-z-colostrum',
        'skora-glowy-i-wlosy', 'genactiv-colostrum-junior-z-czarnym-bzem', 'colostrum-dla-doroslych',
        'colostrum-dla-zwierzat', 'wszystkie-produktu', 'ferie-z-colostrum',
        'new-collection', 'best-selling-collection', 'zestawy-swiateczne',
        'back2school', 'colostrum-a2', 'buduj-odpornosc-dziecka-z-genactiv'
    ]

    coll_data = [['Kolekcja', 'Priorytet']]
    high_priority = ['produkty-dla-dzieci', 'colostrum-dla-doroslych', 'colostrum-dla-zwierzat', 'akcesoria']
    for coll in coll_list:
        priority = 'Wysoki' if coll in high_priority else 'Średni'
        coll_data.append([coll, priority])

    add_styled_table(doc, coll_data, header_color="E67E22")

    doc.add_heading('3.2 Strony bez meta descriptions', level=2)

    pages_data = [
        ['Strona', 'URL', 'Priorytet'],
        ['Kontakt', '/pages/kontakt-3', 'Wysoki'],
        ['Krem dwupak (produkt)', '/products/krem-z-colostrum-genactiv-dwupak', 'Średni'],
    ]
    add_styled_table(doc, pages_data, header_color="E67E22")

    doc.add_heading('3.3 Brakujące Schema.org', level=2)

    schema_missing = [
        ['Typ Schema', 'Strona', 'Korzyść', 'Priorytet'],
        ['FAQPage', '/pages/faq', 'Rich snippets z pytaniami w SERP', 'Wysoki'],
        ['BreadcrumbList', 'Produkty, kolekcje', 'Breadcrumbs w wynikach Google', 'Średni'],
    ]
    add_styled_table(doc, schema_missing, header_color="E67E22")

    doc.add_heading('3.4 Produkty bez recenzji w Schema (28)', level=2)

    no_reviews = doc.add_paragraph()
    no_reviews.add_run('Kategorie produktów bez AggregateRating:\n')

    categories = [
        'Akcesoria (torba, butelka, kosmetyczka) - naturalne, rzadko recenzowane',
        'Dwupaki/multipaki - mogą dziedziczyć recenzje z produktów bazowych',
        'Linia Furever (produkty dla zwierząt) - nowa linia',
        'Fiberbiom - nowy produkt',
        'Produkty Junior w nowych wariantach'
    ]
    for cat in categories:
        doc.add_paragraph(cat, style='List Bullet')

    note = doc.add_paragraph()
    note.add_run('Uwaga: ').bold = True
    note.add_run('Brak recenzji na nowych produktach jest naturalny. ')
    note.add_run('Warto rozważyć kampanię zachęcającą do wystawiania opinii.')

    doc.add_page_break()

    # ===== SECTION 4: TRAFFIC ANALYSIS =====
    doc.add_heading('4. Analiza źródeł ruchu', level=1)

    doc.add_paragraph('Dane z 50 ostatnich zamówień (MCP Shopify, styczeń 2025):')

    traffic_data = [
        ['Źródło', 'Udział', 'Typ', 'Uwagi'],
        ['Google (organic)', '~35%', 'SEO', 'Główne źródło konwersji'],
        ['Instagram', '~25%', 'Social', 'Silny kanał, warto rozbudować UTM'],
        ['Direct', '~20%', 'Direct', 'Powracający klienci, świadomość marki'],
        ['Facebook', '~10%', 'Social', 'Mobile dominant (m.facebook.com)'],
        ['Bing Ads', '~5%', 'Paid', 'UTM tracking działa poprawnie'],
        ['Email/Referral', '~3%', 'Email', 'Poczta Onet - prawdopodobnie newsletter'],
        ['Google Ads', '~2%', 'Paid', 'Niski udział w próbce'],
    ]
    add_styled_table(doc, traffic_data, header_color="0066CC")

    doc.add_paragraph()

    traffic_insights = doc.add_heading('Wnioski:', level=2)

    insights = [
        ('Silna pozycja organiczna', 'Google SEO generuje ~35% zamówień - dobra baza do rozbudowy'),
        ('Potencjał social media', 'Instagram i Facebook łącznie ~35% - warto ujednolicić tracking UTM'),
        ('Paid search', 'Bing Ads ma lepszą widoczność niż Google Ads w próbce - do analizy'),
    ]

    for title, desc in insights:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ===== SECTION 5: RECOMMENDATIONS =====
    doc.add_heading('5. Rekomendacje i priorytety', level=1)

    doc.add_heading('Priorytet 1: Quick wins (1-2 tygodnie)', level=2)

    p1_items = [
        'Uzupełnić meta descriptions dla 16 kolekcji',
        'Dodać meta description do strony kontaktowej',
        'Wdrożyć FAQPage schema na stronie /pages/faq',
        'Ujednolicić UTM tracking dla linków z Instagram'
    ]
    for item in p1_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Priorytet 2: Optymalizacja (2-4 tygodnie)', level=2)

    p2_items = [
        'Wdrożyć BreadcrumbList schema na stronach produktów i kolekcji',
        'Przeanalizować możliwość dziedziczenia recenzji dla wariantów/paków',
        'Zoptymalizować meta descriptions pod kątem CTR (A/B testing)',
        'Rozważyć kampanię zbierania recenzji dla nowych produktów'
    ]
    for item in p2_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Priorytet 3: Rozwój (ongoing)', level=2)

    p3_items = [
        'Kontynuować strategię content marketingu (zgodnie z planem agencji)',
        'Monitorować pozycje long-tail keywords',
        'Wdrożyć 301 redirects z Zooggies (zgodnie z harmonogramem migracji)',
        'Rozbudować content edukacyjny (E-E-A-T)'
    ]
    for item in p3_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ===== SECTION 6: TASK DIVISION =====
    doc.add_heading('6. Podział zadań operacyjnych', level=1)

    intro = doc.add_paragraph()
    intro.add_run('Poniższa tabela przedstawia sugerowany podział zadań między zespół wewnętrzny a agencję SEO, ')
    intro.add_run('uwzględniając kompetencje i dostęp do narzędzi.')

    doc.add_paragraph()

    # ===== INTERNAL TASKS =====
    doc.add_heading('Zadania do realizacji wewnętrznie', level=2)

    internal_note = doc.add_paragraph()
    internal_note.add_run('Poniższe zadania można wykonać bezpośrednio w panelu Shopify Admin ')
    internal_note.add_run('bez konieczności angażowania zewnętrznych zasobów:')

    internal_data = [
        ['Zadanie', 'Gdzie wykonać', 'Szacowany czas'],
        ['Meta descriptions - 16 kolekcji', 'Shopify > Collections > Edit > SEO', '2-3 godziny'],
        ['Meta description - kontakt', 'Shopify > Pages > kontakt-3 > SEO', '10 minut'],
        ['Meta description - produkty (2)', 'Shopify > Products > Edit > SEO', '15 minut'],
        ['UTM linki Instagram', 'Bio + Stories (dodać parametry)', '30 minut'],
        ['Weryfikacja UTM Facebook', 'Business Manager / linki w postach', '30 minut'],
        ['301 redirects Zooggies', 'Shopify > Navigation > URL Redirects', '1-2 godziny'],
        ['Monitorowanie GSC', 'Google Search Console', 'Ongoing'],
    ]
    add_styled_table(doc, internal_data, header_color="27AE60")

    doc.add_paragraph()

    # Instructions box
    doc.add_heading('Instrukcja: Meta descriptions w Shopify', level=3)

    steps = [
        'Zaloguj się do Shopify Admin',
        'Przejdź do Collections (lub Products/Pages)',
        'Wybierz element do edycji',
        'Przewiń do sekcji "Search engine listing preview"',
        'Kliknij "Edit website SEO"',
        'Wpisz meta description (maks. 155 znaków)',
        'Zapisz zmiany'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph()

    # ===== AGENCY TASKS =====
    doc.add_heading('Zadania do realizacji we współpracy z agencją', level=2)

    agency_note = doc.add_paragraph()
    agency_note.add_run('Poniższe zadania wymagają specjalistycznej wiedzy technicznej ')
    agency_note.add_run('lub dostępu do kodu motywu Shopify:')

    agency_data = [
        ['Zadanie', 'Uzasadnienie', 'Priorytet'],
        ['FAQPage Schema', 'Wymaga edycji theme Liquid lub dedykowanego rozwiązania', 'Wysoki'],
        ['BreadcrumbList Schema', 'Wymaga modyfikacji szablonu produktu/kolekcji', 'Średni'],
        ['Optymalizacja Core Web Vitals', 'Wymaga analizy i optymalizacji theme', 'Średni'],
        ['Strategia content SEO', 'Zgodnie z przedstawionym planem rozbudowy treści', 'Wysoki'],
        ['Link building / E-E-A-T', 'Wymaga outreach i relacji z portalami', 'Średni'],
        ['Analiza konkurencji', 'Wymaga narzędzi (Ahrefs, Semrush)', 'Niski'],
    ]
    add_styled_table(doc, agency_data, header_color="3498DB")

    doc.add_paragraph()

    collab_note = doc.add_paragraph()
    collab_note.add_run('Rekomendacja: ').bold = True
    collab_note.add_run('Warto omówić z agencją harmonogram wdrożenia Schema.org (FAQ, Breadcrumbs) ')
    collab_note.add_run('w ramach bieżącej współpracy. Te elementy mogą znacząco poprawić widoczność ')
    collab_note.add_run('w wynikach wyszukiwania (rich snippets).')

    doc.add_page_break()

    # ===== FINAL SUMMARY =====
    doc.add_heading('Podsumowanie', level=1)

    summary = doc.add_paragraph()
    summary.add_run('Stan SEO genactiv.pl należy ocenić jako dobry. ').bold = True
    summary.add_run('Strona posiada solidne fundamenty techniczne - ')
    summary.add_run('prawidłowo wdrożone meta descriptions, Schema.org Product z recenzjami, ')
    summary.add_run('kompletny Open Graph i poprawną strukturę sitemap.\n\n')

    summary.add_run('Zidentyfikowane obszary do poprawy ').bold = True
    summary.add_run('(brakujące meta descriptions dla części kolekcji, Schema FAQ i Breadcrumbs) ')
    summary.add_run('stanowią naturalne kolejne kroki optymalizacji, a nie krytyczne braki.\n\n')

    summary.add_run('Rekomendujemy ').bold = True
    summary.add_run('realizację zadań wewnętrznych (meta descriptions, UTM tracking) równolegle ')
    summary.add_run('z pracami agencji nad rozbudową contentu i wdrożeniem zaawansowanych schematów.')

    # Impact table
    doc.add_paragraph()
    doc.add_heading('Szacowany wpływ rekomendowanych działań:', level=2)

    impact_data = [
        ['Działanie', 'Potencjalny efekt', 'Horyzont'],
        ['Uzupełnienie meta descriptions', '+5-10% CTR na dotkniętych stronach', '2-4 tygodnie'],
        ['FAQPage Schema', 'Rich snippets z FAQ w SERP', '4-8 tygodni'],
        ['BreadcrumbList Schema', 'Breadcrumbs widoczne w Google', '4-8 tygodni'],
        ['Ujednolicenie UTM', 'Lepsza atrybucja, optymalizacja budżetu', 'Natychmiast'],
        ['301 redirects Zooggies', 'Zachowanie equity z ~400 kliknięć/mies.', '4-6 tygodni'],
    ]
    add_styled_table(doc, impact_data, header_color="27AE60")

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    footer = doc.add_paragraph()
    footer.add_run(f'Raport wygenerowany: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n').italic = True
    footer.add_run('Metodologia: Weryfikacja curl + MCP Shopify + analiza sitemap XML\n').italic = True
    footer.add_run('Narzędzia: Claude Code AI Assistant').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_path = '/Users/user/projects/genactiv-klaviyo/seo/GenActiv_SEO_Audyt_Raport_FINAL.docx'
    doc.save(output_path)
    print(f'Raport zapisany: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
