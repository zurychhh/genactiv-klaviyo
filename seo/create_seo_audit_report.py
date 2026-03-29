#!/usr/bin/env python3
"""
Generate SEO Audit Report for GenActiv.pl
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, data, header_color="0066CC", alt_row_color="F5F5F5"):
    """Create a styled table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'

            if i == 0:  # Header row
                set_cell_shading(cell, header_color)
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif i % 2 == 0:  # Alternating rows
                set_cell_shading(cell, alt_row_color)

    return table

def create_report():
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ===== TITLE PAGE =====
    title = doc.add_heading('AUDYT SEO GENACTIV.PL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 102, 204)

    subtitle = doc.add_paragraph('Raport uzupełniony o dane rzeczywiste')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Data audytu: {datetime.now().strftime("%d.%m.%Y")}\n').bold = True
    info.add_run('Źródła danych: WebFetch, MCP Shopify, MCP Klaviyo\n')
    info.add_run('Przygotował: Claude Code AI Assistant')

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading('Spis treści', level=1)
    toc_items = [
        '1. Technical SEO - Stan faktyczny',
        '2. Analiza źródeł ruchu',
        '3. Analiza produktów',
        '4. Porównanie: Rekomendacje vs Rzeczywistość',
        '5. Pilne rekomendacje',
        '6. Szacowany wpływ napraw',
        '7. Podsumowanie audytu',
        '8. Podział zadań: Wewnętrznie vs Agencja'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)

    doc.add_page_break()

    # ===== SECTION 1: TECHNICAL SEO =====
    doc.add_heading('1. Technical SEO - Stan faktyczny', level=1)

    doc.add_paragraph('Analiza kluczowych elementów technicznych SEO na podstawie rzeczywistych danych ze strony genactiv.pl:')

    tech_seo_data = [
        ['Element', 'Status', 'Szczegóły'],
        ['Title tags', '⚠️ Częściowo', 'Produkty OK, homepage niejasny'],
        ['Meta descriptions', '❌ BRAK', 'Nie wykryto na żadnej stronie'],
        ['Schema.org Product', '❌ BRAK', 'Brak JSON-LD dla produktów'],
        ['Schema.org FAQ', '❌ BRAK', 'Brak mimo propozycji FAQ'],
        ['Schema.org Review', '❌ BRAK', 'Brak AggregateRating'],
        ['Open Graph', '❌ BRAK', 'Brak og:title, og:image'],
        ['Canonical URL', '❓ Nieokreślony', 'Wymaga weryfikacji'],
        ['Breadcrumbs', '⚠️ Częściowo', 'Logika OK, brak schema'],
        ['Sitemap.xml', '⚠️ Niekompletny', 'Brak <lastmod> dat'],
        ['Robots.txt', '✅ OK', 'Prawidłowa konfiguracja'],
        ['Mobile viewport', '✅ OK', 'Responsive design OK'],
    ]
    add_styled_table(doc, tech_seo_data)

    doc.add_paragraph()
    critical = doc.add_paragraph()
    critical.add_run('KRYTYCZNE BRAKI TECHNICZNE:').bold = True

    critical_items = [
        'BRAK META DESCRIPTIONS = niski CTR w SERP',
        'BRAK PRODUCT SCHEMA = brak rich snippets z ceną w Google',
        'BRAK REVIEW SCHEMA = brak gwiazdek w wynikach',
        'BRAK OPEN GRAPH = słabe udostępnianie na social media',
        'SITEMAP BEZ LASTMOD = nieefektywne crawlowanie'
    ]
    for item in critical_items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_page_break()

    # ===== SECTION 2: TRAFFIC SOURCES =====
    doc.add_heading('2. Analiza źródeł ruchu', level=1)

    doc.add_paragraph('Dane z 50 ostatnich zamówień (MCP Shopify, styczeń 2025):')

    traffic_data = [
        ['Źródło', 'Udział', 'Typ', 'Uwagi'],
        ['Google SEO', '~35%', 'Organic', 'Główne źródło - dobrze działa'],
        ['Instagram', '~25%', 'Social', 'Silny kanał, brak UTM!'],
        ['Direct', '~20%', 'Direct', 'Powracający + brand'],
        ['Facebook', '~10%', 'Social', 'm.facebook.com - mobile'],
        ['Bing Ads', '~5%', 'Paid', 'UTM tracking działa'],
        ['Poczta Onet', '~3%', 'Referral', 'Newsletter Klaviyo?'],
        ['Google Ads', '~2%', 'Paid', 'Niska obecność'],
    ]
    add_styled_table(doc, traffic_data)

    doc.add_paragraph()
    doc.add_heading('Pozytywne obserwacje:', level=2)
    positives = [
        'UTM tracking działa prawidłowo na kampaniach Bing',
        'Google SEO generuje konwersje (sourceType: "SEO")',
        'Dobry mix kanałów = zdywersyfikowany ruch',
        'Instagram silnie konwertuje'
    ]
    for item in positives:
        p = doc.add_paragraph('✅ ' + item)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_heading('Problemy do naprawy:', level=2)
    problems = [
        'Brak UTM w kampaniach Instagram (tracisz atrybucję)',
        'Niska widoczność Google Ads vs Bing Ads w zamówieniach',
        '"Direct" może zawierać nieoztagowany ruch z kampanii'
    ]
    for item in problems:
        p = doc.add_paragraph('❌ ' + item)
        for run in p.runs:
            run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_page_break()

    # ===== SECTION 3: PRODUCTS =====
    doc.add_heading('3. Analiza produktów (Shopify)', level=1)

    products_data = [
        ['Metryka', 'Wartość', 'Ocena'],
        ['Produkty ACTIVE', '35+', '✅ OK'],
        ['Produkty DRAFT', '~10', '⚠️ Do przeglądu'],
        ['Produkty ARCHIVED', '~8', '✅ Prawidłowo'],
        ['Produkty bez opisu', 'Pojedyncze', '⚠️ Uzupełnić'],
        ['SKU coverage', '~100%', '✅ OK'],
        ['Braki magazynowe', 'MLEKO KLACZY 150g', '⚠️ 0 szt.'],
    ]
    add_styled_table(doc, products_data)

    doc.add_paragraph()
    doc.add_heading('TOP sprzedawane produkty:', level=2)
    top_products = [
        'COLOSTRUM GENACTIV kapsułki (60/120/240 szt.) - bestseller',
        'COLOSTRUM JUNIOR Z CZARNYM BZEM - saszetki i zawiesina',
        'COLOSTRUM GENACTIV proszek 45g - popularna nowość',
        'COLOSTRUM Z CZARNĄ PORZECZKĄ - segment premium',
        'Kosmetyki z Colostrum (maseczka, krem) - cross-sell'
    ]
    for i, item in enumerate(top_products, 1):
        doc.add_paragraph(f'{i}. {item}')

    doc.add_page_break()

    # ===== SECTION 4: COMPARISON =====
    doc.add_heading('4. Porównanie: Rekomendacje vs Rzeczywistość', level=1)

    comparison_data = [
        ['Rekomendacja', 'Odpowiedź agencji', 'Stan na stronie'],
        ['Technical SEO Audit', '0%', '❌ KRYTYCZNE BRAKI'],
        ['Schema.org Product', 'Nie zaplanowano', '❌ BRAK'],
        ['Schema.org FAQ', 'Wspomniano (bez schema)', '❌ BRAK'],
        ['Meta descriptions', 'Nie zaplanowano', '❌ BRAK'],
        ['Core Web Vitals', '0%', '❓ Nie testowano'],
        ['301 Redirects Zooggies', '✅ Zaplanowano', '⏳ Do wdrożenia'],
        ['Content expansion', '✅ ~40% (3 strony)', '⏳ W przygotowaniu'],
        ['Long-tail keywords', '✅ Gap analysis', '⏳ Do wdrożenia'],
        ['CRO (Exit popup, CTA)', '0%', '❌ BRAK'],
        ['Paid-Organic Synergy', '0%', '❌ NIEOPTYMALNIE'],
    ]
    add_styled_table(doc, comparison_data)

    doc.add_page_break()

    # ===== SECTION 5: RECOMMENDATIONS =====
    doc.add_heading('5. Pilne rekomendacje', level=1)

    # Priority 1
    p1 = doc.add_heading('🔴 PRIORYTET 1: Technical SEO (natychmiast)', level=2)
    priority1 = [
        'Dodać meta descriptions do WSZYSTKICH stron produktowych',
        'Wdrożyć Schema.org Product JSON-LD (cena, dostępność, SKU)',
        'Dodać Schema.org AggregateRating (jeśli są recenzje)',
        'Dodać Open Graph tags (og:title, og:image, og:description)',
        'Uzupełnić sitemap.xml o <lastmod> daty',
        'Zweryfikować canonical URLs'
    ]
    for item in priority1:
        doc.add_paragraph('☐ ' + item, style='List Bullet')

    # Priority 2
    doc.add_heading('🟠 PRIORYTET 2: Zooggies Migration', level=2)
    priority2 = [
        'Wdrożyć mapowanie 301 redirects (arkusz gotowy)',
        'Przenieść content z TOP10 fraz',
        'Monitorować GSC po migracji'
    ]
    for item in priority2:
        doc.add_paragraph('☐ ' + item, style='List Bullet')

    # Priority 3
    doc.add_heading('🟡 PRIORYTET 3: Content & Keywords', level=2)
    priority3 = [
        'Opublikować 3 strony produktowe (Colostrum/Fiberbiom/Junior)',
        'Wdrożyć FAQ z Schema.org markup',
        'Rozbudować long-tail content (gap analysis)'
    ]
    for item in priority3:
        doc.add_paragraph('☐ ' + item, style='List Bullet')

    # Priority 4
    doc.add_heading('🟢 PRIORYTET 4: Tracking & Attribution', level=2)
    priority4 = [
        'Dodać UTM do WSZYSTKICH linków Instagram',
        'Zweryfikować Google Ads tracking',
        'Ustawić cele konwersji w GA4'
    ]
    for item in priority4:
        doc.add_paragraph('☐ ' + item, style='List Bullet')

    doc.add_page_break()

    # ===== SECTION 6: IMPACT =====
    doc.add_heading('6. Szacowany wpływ napraw', level=1)

    impact_data = [
        ['Naprawa', 'Szacowany wzrost', 'Timeframe'],
        ['Meta descriptions', '+15-25% CTR', '2-4 tyg.'],
        ['Product Schema', '+10-20% CTR (rich snippets)', '2-6 tyg.'],
        ['Review Schema', '+5-15% CTR (gwiazdki)', '2-6 tyg.'],
        ['Zooggies 301s', 'Zachowanie ~400 kliknięć/mies.', 'Natychmiast'],
        ['Open Graph', 'Lepsze social shares', 'Natychmiast'],
        ['Long-tail content', '+30-50% organic traffic', '3-6 mies.'],
    ]
    add_styled_table(doc, impact_data, header_color="27AE60")

    doc.add_page_break()

    # ===== SECTION 7: SUMMARY =====
    doc.add_heading('7. Podsumowanie audytu', level=1)

    summary_title = doc.add_paragraph()
    summary_title.add_run('Stan SEO genactiv.pl: 4/10 ').bold = True
    summary_title.add_run('⚠️').font.size = Pt(16)

    summary_data = [
        ['Obszar', 'Ocena', 'Komentarz'],
        ['Ruch organiczny', '7/10', 'Działa, Google SEO konwertuje'],
        ['Technical SEO', '2/10', 'Krytyczne braki w structured data'],
        ['On-page SEO', '4/10', 'Brak meta descriptions'],
        ['Content', '5/10', 'Opisy OK, brak rozbudowanego contentu'],
        ['Mobile', '7/10', 'Responsive działa'],
        ['Tracking', '5/10', 'UTM na Bing OK, Instagram bez'],
    ]
    add_styled_table(doc, summary_data, header_color="1A3B5D")

    doc.add_paragraph()

    doc.add_heading('Co działa:', level=2)
    works = [
        'Google SEO generuje sprzedaż (~35% zamówień)',
        'Shopify działa stabilnie',
        'UTM tracking na kampaniach Bing',
        'Podstawowa responsywność'
    ]
    for item in works:
        p = doc.add_paragraph('✅ ' + item)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_heading('Co nie działa:', level=2)
    not_works = [
        'Brak meta descriptions (100% stron)',
        'Brak Schema.org dla e-commerce',
        'Brak Open Graph',
        'Sitemap bez dat modyfikacji',
        'Instagram bez UTM tracking'
    ]
    for item in not_works:
        p = doc.add_paragraph('❌ ' + item)
        for run in p.runs:
            run.font.color.rgb = RGBColor(200, 0, 0)

    doc.add_page_break()

    # ===== SECTION 8: TASK DIVISION =====
    doc.add_heading('8. Podział zadań: Wewnętrznie vs Agencja', level=1)

    intro = doc.add_paragraph()
    intro.add_run('Poniższa tabela pokazuje, które zadania można wykonać wewnętrznie (mając dostęp do Shopify, Klaviyo, Google Search Console) a które wymagają zaangażowania agencji SEO.').italic = True

    doc.add_paragraph()

    # ===== INTERNAL TASKS =====
    doc.add_heading('✅ ZADANIA DO WYKONANIA WEWNĘTRZNIE', level=2)

    internal_data = [
        ['Zadanie', 'Narzędzie', 'Jak to zrobić'],
        ['Meta descriptions produktów', 'Shopify Admin', 'Products → Edit → SEO → Meta description'],
        ['Meta descriptions kolekcji', 'Shopify Admin', 'Collections → Edit → SEO → Meta description'],
        ['Open Graph images', 'Shopify Admin', 'Settings → Files → Upload + Product images'],
        ['UTM tagi Instagram', 'Link builder', 'Dodać ?utm_source=instagram&utm_medium=social do wszystkich linków w bio i stories'],
        ['UTM tagi Facebook', 'Link builder', 'Analogicznie jak Instagram'],
        ['Weryfikacja GA4', 'Google Analytics', 'Admin → Data Streams → Sprawdzić eventy'],
        ['301 redirects Zooggies', 'Shopify Admin', 'Online Store → Navigation → URL Redirects → Import CSV'],
        ['Opisy produktów', 'Shopify Admin', 'Products → Edit → Description (HTML)'],
        ['Alt tagi obrazów', 'Shopify Admin', 'Products → Media → Edit alt text'],
        ['Sitemap submission', 'Google Search Console', 'Sitemaps → Add sitemap.xml'],
        ['Monitoring GSC', 'Google Search Console', 'Performance → sprawdzać CTR, pozycje'],
        ['Cele konwersji GA4', 'Google Analytics', 'Admin → Conversions → New event'],
        ['Email UTM tracking', 'Klaviyo', 'Campaigns → Settings → UTM tracking ON'],
    ]
    add_styled_table(doc, internal_data, header_color="27AE60")

    doc.add_paragraph()

    # Detailed instructions
    doc.add_heading('Szczegółowe instrukcje dla kluczowych zadań:', level=3)

    # Meta descriptions
    doc.add_paragraph()
    meta_title = doc.add_paragraph()
    meta_title.add_run('1. Meta descriptions w Shopify:').bold = True
    meta_steps = [
        'Zaloguj się do Shopify Admin (genactiv.myshopify.com/admin)',
        'Przejdź do Products → All products',
        'Kliknij produkt do edycji',
        'Przewiń do sekcji "Search engine listing preview"',
        'Kliknij "Edit website SEO"',
        'Wpisz meta description (max 155 znaków, zawierający słowa kluczowe)',
        'Zapisz zmiany',
        'Powtórz dla wszystkich ~35 aktywnych produktów'
    ]
    for i, step in enumerate(meta_steps, 1):
        doc.add_paragraph(f'   {i}. {step}')

    # 301 redirects
    doc.add_paragraph()
    redirect_title = doc.add_paragraph()
    redirect_title.add_run('2. 301 Redirects dla Zooggies:').bold = True
    redirect_steps = [
        'Przygotuj plik CSV z kolumnami: Redirect from, Redirect to',
        'Format: /stary-url, /nowy-url (bez domeny)',
        'W Shopify Admin → Online Store → Navigation',
        'Kliknij "URL Redirects" → "Import"',
        'Wgraj plik CSV',
        'Zweryfikuj przekierowania testując stare URLe'
    ]
    for i, step in enumerate(redirect_steps, 1):
        doc.add_paragraph(f'   {i}. {step}')

    # UTM Instagram
    doc.add_paragraph()
    utm_title = doc.add_paragraph()
    utm_title.add_run('3. UTM dla Instagram:').bold = True
    doc.add_paragraph('   Każdy link do genactiv.pl powinien wyglądać tak:')
    utm_example = doc.add_paragraph('   https://genactiv.pl/products/colostrum-genactiv-120-kapsulek?utm_source=instagram&utm_medium=social&utm_campaign=bio_link')
    utm_example.runs[0].font.name = 'Courier New'
    utm_example.runs[0].font.size = Pt(9)

    doc.add_page_break()

    # ===== AGENCY TASKS =====
    doc.add_heading('🔶 ZADANIA WYMAGAJĄCE AGENCJI', level=2)

    agency_data = [
        ['Zadanie', 'Powód', 'Alternatywa wewnętrzna'],
        ['Schema.org Product JSON-LD', 'Wymaga edycji theme Liquid', 'Można: Shopify app "JSON-LD for SEO"'],
        ['Schema.org FAQ', 'Wymaga edycji theme Liquid', 'Można: Shopify app lub manual JSON-LD'],
        ['Schema.org Review/Rating', 'Wymaga integracji z reviews', 'Można: Judge.me, Loox (apps)'],
        ['Sitemap <lastmod>', 'Shopify generuje automatycznie', 'Nie można zmienić w standard Shopify'],
        ['Core Web Vitals optymalizacja', 'Wymaga optymalizacji theme', 'Częściowo: kompresja obrazów'],
        ['Rozbudowany content SEO', 'Wymaga copywritera', 'Można: AI-assisted (Claude)'],
        ['Link building zewnętrzny', 'Wymaga outreach', 'Można: portale z listy agencji'],
        ['Audyt konkurencji', 'Wymaga narzędzi (Ahrefs, Semrush)', 'Można: darmowe narzędzia (ograniczone)'],
    ]
    add_styled_table(doc, agency_data, header_color="E67E22")

    doc.add_paragraph()

    # Apps recommendation
    doc.add_heading('Rekomendowane aplikacje Shopify (alternatywa dla agencji):', level=3)

    apps = [
        ('JSON-LD for SEO by Ilana Davis', 'Automatyczny Schema.org', 'Free / $299 one-time'),
        ('Smart SEO', 'Meta tags, JSON-LD, sitemap', '$9.99/mies.'),
        ('SEO Manager', 'Kompleksowe SEO', '$20/mies.'),
        ('Judge.me', 'Reviews + Schema', 'Free / $15/mies.'),
        ('TinyIMG', 'Kompresja obrazów (Core Web Vitals)', 'Free / $9.99/mies.'),
    ]

    apps_data = [['Aplikacja', 'Funkcja', 'Cena']] + [list(app) for app in apps]
    add_styled_table(doc, apps_data, header_color="9B59B6")

    doc.add_page_break()

    # ===== FINAL SUMMARY TABLE =====
    doc.add_heading('📋 PODSUMOWANIE: CO ROBIĆ SAMEMU vs AGENCJA', level=2)

    final_data = [
        ['Zadanie', 'Kto', 'Priorytet', 'Czas realizacji'],
        ['Meta descriptions (35 produktów)', 'WEWNĘTRZNIE', '🔴 Pilne', '1-2 dni'],
        ['301 redirects Zooggies', 'WEWNĘTRZNIE', '🔴 Pilne', '1 dzień'],
        ['UTM tracking Instagram/FB', 'WEWNĘTRZNIE', '🟠 Wysoki', '2-3 godz.'],
        ['Alt tagi obrazów', 'WEWNĘTRZNIE', '🟡 Średni', '1 dzień'],
        ['Open Graph (podstawowy)', 'WEWNĘTRZNIE', '🟡 Średni', '2-3 godz.'],
        ['Schema.org Product', 'AGENCJA lub APP', '🔴 Pilne', '1-2 tyg.'],
        ['Schema.org FAQ', 'AGENCJA lub APP', '🟠 Wysoki', '1 tydz.'],
        ['Core Web Vitals', 'AGENCJA', '🟡 Średni', '2-4 tyg.'],
        ['Content rozbudowany', 'AGENCJA', '🟡 Średni', '4-8 tyg.'],
        ['Link building', 'AGENCJA', '🟢 Niski', 'Ongoing'],
    ]
    add_styled_table(doc, final_data, header_color="1A3B5D")

    doc.add_paragraph()

    # Final note
    final_note = doc.add_paragraph()
    final_note.add_run('REKOMENDACJA KOŃCOWA: ').bold = True
    final_note.add_run('Rozpocznij od zadań wewnętrznych (meta descriptions, 301 redirects, UTM tracking) - to 60% sukcesu przy 0 zł kosztów. Dla Schema.org rozważ aplikację Shopify zamiast agencji - szybciej i taniej.')

    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    footer = doc.add_paragraph()
    footer.add_run(f'Raport wygenerowany: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n').italic = True
    footer.add_run('Narzędzia: Claude Code + MCP (Shopify, Klaviyo, WebFetch)').italic = True
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save document
    output_path = '/Users/user/projects/genactiv-klaviyo/seo/GenActiv_SEO_Audyt_Raport_2025.docx'
    doc.save(output_path)
    print(f'✅ Raport zapisany: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
