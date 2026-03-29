#!/usr/bin/env python3
"""
Generuje zaktualizowany raport SEO z listą możliwości Claude Code
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Ustaw kolor tła komórki"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_report():
    doc = Document()

    # Tytuł
    title = doc.add_heading('GenActiv.pl - Raport SEO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Możliwości automatyzacji Claude Code')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f'Data: {datetime.now().strftime("%d.%m.%Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # SEKCJA 1: Podsumowanie testów
    doc.add_heading('1. Wyniki testów możliwości', level=1)

    p = doc.add_paragraph()
    p.add_run('Przeprowadzono kompleksowe testy wszystkich funkcjonalności SEO. ')
    p.add_run('Wszystkie testy zakończyły się sukcesem.').bold = True

    # Tabela testów
    test_table = doc.add_table(rows=5, cols=3)
    test_table.style = 'Table Grid'

    headers = ['Test', 'Funkcjonalność', 'Status']
    for i, header in enumerate(headers):
        cell = test_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E8E8E8')

    tests = [
        ('Test 1', 'Update Product SEO via Shopify API', '✅ DZIAŁA'),
        ('Test 2', 'Update Collection SEO via Shopify API', '✅ DZIAŁA'),
        ('Test 3', 'Shopify Theme API - odczyt plików Liquid', '✅ DZIAŁA'),
        ('Test 4', 'Shopify Theme API - zapis plików Liquid', '✅ DZIAŁA'),
    ]

    for i, (test, func, status) in enumerate(tests, 1):
        test_table.rows[i].cells[0].text = test
        test_table.rows[i].cells[1].text = func
        test_table.rows[i].cells[2].text = status

    doc.add_paragraph()

    # SEKCJA 2: Co mogę zrobić sam
    doc.add_heading('2. Co Claude Code może zrobić samodzielnie', level=1)

    # 2.1 Meta descriptions
    doc.add_heading('2.1 Uzupełnienie meta descriptions kolekcji', level=2)

    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('28 kolekcji nie ma meta descriptions.')

    p = doc.add_paragraph()
    p.add_run('Rozwiązanie: ').bold = True
    p.add_run('Mogę automatycznie wygenerować i dodać meta descriptions dla wszystkich kolekcji poprzez Shopify GraphQL API.')

    p = doc.add_paragraph()
    p.add_run('Jak to zrobię:\n').bold = True

    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Pobieram listę kolekcji bez meta description')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Generuję optymalne meta descriptions (max 160 znaków) na podstawie nazwy i produktów')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Zapisuję przez API: mutation collectionUpdate')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Weryfikuję zmiany na stronie')

    # Lista kolekcji do poprawy
    p = doc.add_paragraph()
    p.add_run('Kolekcje wymagające meta descriptions:').bold = True

    collections = [
        'Nowości', 'Najlepiej sprzedające się produkty', 'Zestawy Świąteczne',
        'Colostrum i Mleko Klaczy', 'Colostrum tabletki do ssania', 'Colostrum proszek',
        'Colostrum kapsułki', 'COLOSTRUM JUNIOR CZARNY BEZ', 'Colostrum GENACTIV',
        'Linia COLOSTRUM GENACTIV', 'Colostrum z czystej siary', 'Suplement dla dwojga',
        'Odporność', 'Kolostrum COLOSTRIGEN', 'Regeneracja tkanek', 'Black Friday 2024',
        'Alergia i nietolerancja pokarmowa', 'Colosregen', 'Cukrzyca', 'Dermatologia',
        'Dieta i odchudzanie', 'Układ nerwowy', 'Układ oddechowy', 'Uroda i włosy',
        'Colostrum dla sportowców', 'Układ mięśniowo-szkieletowy'
    ]

    for coll in collections[:15]:
        doc.add_paragraph(f'• {coll}', style='List Bullet')
    doc.add_paragraph(f'... i {len(collections)-15} więcej')

    # 2.2 FAQPage Schema
    doc.add_heading('2.2 Dodanie FAQPage Schema', level=2)

    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('Strona FAQ nie ma struktury FAQPage Schema.org - brak rich snippets w Google.')

    p = doc.add_paragraph()
    p.add_run('Rozwiązanie: ').bold = True
    p.add_run('Dodam FAQPage Schema do pliku templates/page.faq.json lub snippets.')

    p = doc.add_paragraph()
    p.add_run('Jak to zrobię:\n').bold = True

    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Pobieram aktualny plik FAQ template')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Parseuję pytania i odpowiedzi z sekcji FAQ')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Generuję JSON-LD z FAQPage Schema')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Dodaję snippet do theme.liquid lub tworzę nowy snippet')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Testuję w Google Rich Results Test')

    # Przykład kodu
    p = doc.add_paragraph()
    p.add_run('Przykład docelowego Schema:').bold = True

    code = doc.add_paragraph()
    code.add_run('''<script type="application/ld+json">
{
  "@context": "https://schema.org",
  "@type": "FAQPage",
  "mainEntity": [
    {
      "@type": "Question",
      "name": "W jakim czasie od rozpoczęcia laktacji...",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "Colostrum wytwarzane jest..."
      }
    }
  ]
}
</script>''').font.size = Pt(9)

    # 2.3 BreadcrumbList Schema
    doc.add_heading('2.3 Dodanie BreadcrumbList Schema', level=2)

    p = doc.add_paragraph()
    p.add_run('Problem: ').bold = True
    p.add_run('Brak BreadcrumbList Schema na stronach produktów i kolekcji.')

    p = doc.add_paragraph()
    p.add_run('Rozwiązanie: ').bold = True
    p.add_run('Dodam BreadcrumbList Schema do szablonów product i collection.')

    p = doc.add_paragraph()
    p.add_run('Jak to zrobię:\n').bold = True

    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Stworzę snippet snippets/breadcrumb-schema.liquid')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Wykorzystam zmienne Liquid: product.title, collection.title, shop.name')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Dołączę snippet do sections/main-product.liquid i main-collection.liquid')
    steps = doc.add_paragraph(style='List Bullet')
    steps.add_run('Zweryfikuję działanie')

    doc.add_paragraph()

    # SEKCJA 3: Tabela podsumowująca
    doc.add_heading('3. Podsumowanie - tabela zadań', level=1)

    task_table = doc.add_table(rows=8, cols=4)
    task_table.style = 'Table Grid'

    headers = ['Zadanie', 'Kto wykonuje', 'Metoda', 'Priorytet']
    for i, header in enumerate(headers):
        cell = task_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '0066CC')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tasks = [
        ('Meta descriptions - 28 kolekcji', 'Claude Code', 'Shopify GraphQL API\nmutation collectionUpdate', 'WYSOKI'),
        ('Meta descriptions - strona kontakt', 'Claude Code', 'Shopify GraphQL API\nmutation pageUpdate', 'ŚREDNI'),
        ('FAQPage Schema na /faq', 'Claude Code', 'Shopify Theme API\nsnippets/faq-schema.liquid', 'WYSOKI'),
        ('BreadcrumbList Schema', 'Claude Code', 'Shopify Theme API\nsnippets/breadcrumb-schema.liquid', 'ŚREDNI'),
        ('Review Schema (gwiazdki)', 'Już działa ✅', 'Automatyczne z AggregateRating', '-'),
        ('Product Schema', 'Już działa ✅', 'Automatyczne w theme', '-'),
        ('301 redirecty Zooggies', 'Agencja/Admin', 'Shopify Admin → Navigation → URL Redirects', 'NISKI'),
    ]

    for i, (task, who, method, priority) in enumerate(tasks, 1):
        task_table.rows[i].cells[0].text = task
        task_table.rows[i].cells[1].text = who
        task_table.rows[i].cells[2].text = method
        task_table.rows[i].cells[3].text = priority

        # Kolorowanie "Claude Code"
        if who == 'Claude Code':
            set_cell_shading(task_table.rows[i].cells[1], 'E8F5E9')
        elif 'Już działa' in who:
            set_cell_shading(task_table.rows[i].cells[1], 'E3F2FD')

    doc.add_paragraph()

    # SEKCJA 4: Plan działania
    doc.add_heading('4. Proponowany plan działania', level=1)

    p = doc.add_paragraph()
    p.add_run('Kolejność wykonania zadań przez Claude Code:').bold = True

    plan = [
        ('Faza 1', 'Uzupełnienie meta descriptions dla 28 kolekcji', 'Największy wpływ na SEO - poprawa widoczności kolekcji w Google'),
        ('Faza 2', 'Dodanie FAQPage Schema', 'Możliwość wyświetlania FAQ w wynikach Google jako rich snippets'),
        ('Faza 3', 'Dodanie BreadcrumbList Schema', 'Lepsza struktura nawigacyjna w wynikach wyszukiwania'),
        ('Faza 4', 'Meta description dla strony kontakt', 'Drobna poprawa'),
    ]

    for phase, task, reason in plan:
        p = doc.add_paragraph()
        p.add_run(f'{phase}: ').bold = True
        p.add_run(f'{task}\n')
        p.add_run(f'   Powód: {reason}').italic = True

    doc.add_paragraph()

    # SEKCJA 5: Wnioski
    doc.add_heading('5. Wnioski', level=1)

    p = doc.add_paragraph()
    p.add_run('Kluczowe ustalenia:\n').bold = True

    conclusions = [
        'Claude Code ma pełny dostęp do Shopify API (GraphQL + Theme API) i może samodzielnie wykonać większość zadań SEO',
        'Meta descriptions, FAQPage Schema i BreadcrumbList Schema mogą być dodane automatycznie',
        'Product Schema, Review Schema i AggregateRating już działają poprawnie na stronie',
        'Nie wymaga to dodatkowej pracy agencji - wystarczy akceptacja planu i uruchomienie automatyzacji'
    ]

    for conc in conclusions:
        doc.add_paragraph(f'• {conc}', style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('\nNastępny krok: ').bold = True
    p.add_run('Po akceptacji tego planu, Claude Code może natychmiast rozpocząć implementację zaczynając od meta descriptions dla kolekcji.')

    # Zapisz
    output_path = '/Users/user/projects/genactiv-klaviyo/seo/GenActiv_SEO_Mozliwosci_Claude.docx'
    doc.save(output_path)
    print(f"Raport zapisany: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
