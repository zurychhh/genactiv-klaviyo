#!/usr/bin/env python3
"""
Tworzy zaktualizowany dokument Word ze wszystkimi poprawionymi meta descriptions
Poprawki: Genactiv (nie GenActiv), bez napojów/koktajli, maseczki do twarzy
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Ustaw kolor tła komórki"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_report():
    doc = Document()

    # Tytuł
    title = doc.add_heading('GenActiv.pl - Specyfikacja zmian SEO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Pełna lista zmian do wdrożenia')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f'Data: {datetime.now().strftime("%d.%m.%Y")} | Status: WDROŻONE')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ===========================================
    # SEKCJA 1: META DESCRIPTIONS KOLEKCJI
    # ===========================================
    doc.add_heading('1. Meta Descriptions - Kolekcje (24 szt.) ✅ WDROŻONE', level=1)

    p = doc.add_paragraph()
    p.add_run('Status: ').bold = True
    p.add_run('Wszystkie 24 kolekcje zostały zaktualizowane przez Shopify GraphQL API.')

    p = doc.add_paragraph()
    p.add_run('Poprawki wprowadzone:').bold = True

    doc.add_paragraph('• Genactiv (nie GenActiv) - poprawna pisownia marki', style='List Bullet')
    doc.add_paragraph('• Usunięto "do napojów i koktajli" z Colostrum proszek', style='List Bullet')
    doc.add_paragraph('• Maseczki z Colostrum - do skóry twarzy (nie włosów)', style='List Bullet')

    doc.add_paragraph()

    # Tabela z meta descriptions
    collections = [
        ("Nowości", "Odkryj najnowsze produkty Genactiv z colostrum bovinum. Sprawdź nowości wspierające odporność, regenerację i piękną skórę."),
        ("Najlepiej sprzedające się produkty", "Najpopularniejsze produkty Genactiv z colostrum bovinum. Zobacz bestsellery wybierane przez tysiące zadowolonych klientów."),
        ("Zestawy Świąteczne", "Świąteczne zestawy prezentowe Genactiv z colostrum. Podaruj zdrowie i odporność bliskim w wyjątkowej cenie."),
        ("Colostrum i Mleko Klaczy", "Colostrum bovinum i mleko klaczy Genactiv. Naturalne wsparcie odporności i regeneracji dla całej rodziny."),
        ("Colostrum tabletki do ssania", "Colostrum Genactiv w tabletkach do ssania. Wygodna forma dla dzieci i dorosłych - smaczne i skuteczne wsparcie odporności."),
        ("Colostrum proszek", "Colostrum Genactiv w proszku. Najwyższa biodostępność składników aktywnych - idealne dla wymagających."),
        ("Colostrum kapsułki", "Colostrum Genactiv w kapsułkach. Wygodna codzienna suplementacja dla osób dorosłych wspierająca odporność."),
        ("COLOSTRUM JUNIOR CZARNY BEZ", "Colostrum Junior z czarnym bzem Genactiv. Pyszne tabletki do ssania dla dzieci wzmacniające odporność."),
        ("Back2school", "Przygotuj dziecko na powrót do szkoły z Genactiv Colostrum. Wspieraj odporność malucha naturalnym colostrum."),
        ("Akcesoria", "Akcesoria Genactiv – praktyczne dodatki do suplementacji colostrum. Miarki, pojemniki i inne przydatne produkty."),
        ("Maseczki z Colostrum", "Maseczki do twarzy z colostrum Genactiv. Intensywna regeneracja i nawilżenie skóry twarzy dzięki naturalnym składnikom."),
        ("Kremy z Colostrum", "Kremy z colostrum Genactiv dla pięknej skóry. Naturalna pielęgnacja wspierająca regenerację i nawilżenie."),
        ("Skóra głowy i włosy", "Kosmetyki Genactiv z colostrum do pielęgnacji skóry głowy i włosów. Naturalne wzmocnienie i regeneracja."),
        ("Genactiv Colostrum Junior z Czarnym Bzem", "Colostrum Junior z czarnym bzem - pyszne tabletki do ssania dla dzieci. Naturalne wsparcie odporności od Genactiv."),
        ("Colostrum dla dzieci", "Colostrum Genactiv dla dzieci - naturalne wsparcie odporności maluchów. Bezpieczne produkty w smacznych formach."),
        ("Colostrum dla dorosłych", "Colostrum Genactiv dla dorosłych. Kapsułki, tabletki i proszek wspierające odporność i regenerację organizmu."),
        ("Buduj odporność dziecka", "Wzmocnij odporność dziecka naturalnym colostrum Genactiv. Sprawdzone produkty dla najmłodszych."),
        ("Colostrum A2", "Colostrum A2 Genactiv z mleka krów A2. Premium colostrum dla osób wrażliwych na białko A1."),
        ("Colostrum dla zwierząt", "Colostrum Genactiv dla zwierząt. Naturalne wsparcie odporności i zdrowia Twojego pupila."),
        ("Colostrum dla psów", "Colostrum Genactiv dla psów. Wspieraj zdrowie i odporność swojego psa naturalnym colostrum bovinum."),
        ("Colostrum dla kotów", "Colostrum Genactiv dla kotów. Naturalne wsparcie odporności i zdrowia Twojego kota."),
        ("Colostrum dla koni", "Colostrum Genactiv dla koni. Profesjonalne wsparcie odporności i regeneracji dla koni sportowych i hodowlanych."),
        ("Wszystkie produkty", "Pełna oferta produktów Genactiv z colostrum bovinum. Suplementy, kosmetyki i produkty dla zwierząt."),
        ("-20% z kodem FERIE", "Promocja zimowa Genactiv! Odbierz 20% rabatu na produkty z colostrum z kodem FERIE. Wspieraj odporność w niższej cenie."),
    ]

    # Tabela
    table = doc.add_table(rows=len(collections)+1, cols=2)
    table.style = 'Table Grid'

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Kolekcja'
    header_cells[1].text = 'Meta Description'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '0066CC')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, (name, meta) in enumerate(collections, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = meta
        # Zielone tło dla wdrożonych
        set_cell_shading(table.rows[i].cells[0], 'E8F5E9')

    doc.add_paragraph()

    # ===========================================
    # SEKCJA 2: FAQPage Schema
    # ===========================================
    doc.add_heading('2. FAQPage Schema - Do wdrożenia', level=1)

    p = doc.add_paragraph()
    p.add_run('Cel: ').bold = True
    p.add_run('Dodanie struktury FAQPage Schema.org do strony /pages/faq')

    p = doc.add_paragraph()
    p.add_run('Lokalizacja pliku: ').bold = True
    p.add_run('snippets/faq-schema.liquid')

    p = doc.add_paragraph()
    p.add_run('Kod do dodania:').bold = True

    code = doc.add_paragraph()
    schema_code = '''{% if template.suffix == 'faq' or request.path == '/pages/faq' %}
<script type="application/ld+json">
{
  "@context": "https://schema.org",
  "@type": "FAQPage",
  "mainEntity": [
    {
      "@type": "Question",
      "name": "W jakim czasie od rozpoczęcia laktacji pozyskuje się colostrum?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "Colostrum wytwarzane jest przez krowy w pierwszych 24-72 godzinach po porodzie. Genactiv pozyskuje colostrum wyłącznie z pierwszego udoju, które zawiera najwyższe stężenie immunoglobulin i czynników wzrostu."
      }
    },
    {
      "@type": "Question",
      "name": "Czy colostrum jest bezpieczne dla osób z nietolerancją laktozy?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "Colostrum zawiera minimalną ilość laktozy. Większość osób z łagodną nietolerancją laktozy dobrze toleruje colostrum. W przypadku silnej nietolerancji zalecamy konsultację z lekarzem."
      }
    },
    {
      "@type": "Question",
      "name": "Jak przechowywać produkty z colostrum?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "Produkty Genactiv należy przechowywać w suchym, chłodnym miejscu, z dala od światła słonecznego. Po otwarciu opakowania proszku należy zużyć produkt w ciągu 30 dni."
      }
    },
    {
      "@type": "Question",
      "name": "Od jakiego wieku można podawać colostrum dzieciom?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "Colostrum Junior Genactiv jest przeznaczone dla dzieci od 3. roku życia. Dla młodszych dzieci zalecamy konsultację z pediatrą."
      }
    },
    {
      "@type": "Question",
      "name": "Czy colostrum można łączyć z innymi suplementami?",
      "acceptedAnswer": {
        "@type": "Answer",
        "text": "Tak, colostrum można bezpiecznie łączyć z większością suplementów diety. Zalecamy przyjmowanie colostrum na czczo, 30 minut przed posiłkiem dla najlepszego wchłaniania."
      }
    }
  ]
}
</script>
{% endif %}'''
    code.add_run(schema_code).font.size = Pt(8)

    doc.add_paragraph()

    # ===========================================
    # SEKCJA 3: BreadcrumbList Schema
    # ===========================================
    doc.add_heading('3. BreadcrumbList Schema - Do wdrożenia', level=1)

    p = doc.add_paragraph()
    p.add_run('Cel: ').bold = True
    p.add_run('Dodanie nawigacji okruszkowej w strukturze Schema.org')

    p = doc.add_paragraph()
    p.add_run('Lokalizacja pliku: ').bold = True
    p.add_run('snippets/breadcrumb-schema.liquid')

    p = doc.add_paragraph()
    p.add_run('Kod do dodania:').bold = True

    code = doc.add_paragraph()
    breadcrumb_code = '''<script type="application/ld+json">
{
  "@context": "https://schema.org",
  "@type": "BreadcrumbList",
  "itemListElement": [
    {
      "@type": "ListItem",
      "position": 1,
      "name": "Strona główna",
      "item": "{{ shop.url }}"
    }
    {%- if template.name == 'collection' -%}
    ,{
      "@type": "ListItem",
      "position": 2,
      "name": "{{ collection.title }}",
      "item": "{{ shop.url }}{{ collection.url }}"
    }
    {%- endif -%}
    {%- if template.name == 'product' -%}
    {%- if product.collections.first -%}
    ,{
      "@type": "ListItem",
      "position": 2,
      "name": "{{ product.collections.first.title }}",
      "item": "{{ shop.url }}{{ product.collections.first.url }}"
    }
    {%- endif -%}
    ,{
      "@type": "ListItem",
      "position": {% if product.collections.first %}3{% else %}2{% endif %},
      "name": "{{ product.title }}",
      "item": "{{ shop.url }}{{ product.url }}"
    }
    {%- endif -%}
  ]
}
</script>'''
    code.add_run(breadcrumb_code).font.size = Pt(8)

    p = doc.add_paragraph()
    p.add_run('\nDołączenie w theme.liquid: ').bold = True
    p.add_run("{% render 'breadcrumb-schema' %}")

    doc.add_paragraph()

    # ===========================================
    # SEKCJA 4: Podsumowanie
    # ===========================================
    doc.add_heading('4. Podsumowanie wdrożenia', level=1)

    summary_table = doc.add_table(rows=4, cols=3)
    summary_table.style = 'Table Grid'

    # Header
    headers = ['Zadanie', 'Status', 'Metoda']
    for i, header in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '0066CC')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Rows
    tasks = [
        ('Meta descriptions - 24 kolekcje', '✅ WDROŻONE', 'Shopify GraphQL API'),
        ('FAQPage Schema', '⏳ Do wdrożenia', 'Shopify Theme API'),
        ('BreadcrumbList Schema', '⏳ Do wdrożenia', 'Shopify Theme API'),
    ]

    for i, (task, status, method) in enumerate(tasks, 1):
        summary_table.rows[i].cells[0].text = task
        summary_table.rows[i].cells[1].text = status
        summary_table.rows[i].cells[2].text = method

        if '✅' in status:
            set_cell_shading(summary_table.rows[i].cells[1], 'E8F5E9')
        else:
            set_cell_shading(summary_table.rows[i].cells[1], 'FFF3E0')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Uwagi:').bold = True
    doc.add_paragraph('• Product Schema i AggregateRating już działają poprawnie na stronie', style='List Bullet')
    doc.add_paragraph('• FAQPage i BreadcrumbList Schema wymagają edycji plików motywu Shopify', style='List Bullet')
    doc.add_paragraph('• Wszystkie meta descriptions używają poprawnej pisowni: Genactiv', style='List Bullet')

    # Zapisz
    output_path = '/Users/user/projects/genactiv-klaviyo/seo/SEO_SPECYFIKACJA_ZMIAN_FINAL.docx'
    doc.save(output_path)
    print(f"Dokument zapisany: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
